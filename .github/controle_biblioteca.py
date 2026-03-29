#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
SISTEMA DE CONTROLE DE ESTOQUE DE LIVROS - BIBLIOTECA ESCOLAR
Escola Estadual de Educação Básica Nicolau Müssnich
Rua Professora Nely Muller 72 - Boa União - Estrela/RS

Desenvolvido em Python com Tkinter
Formato de data: DD/MM/YYYY (Brasileiro)
Exportação: Word (.docx)
"""

import tkinter as tk
from tkinter import ttk, messagebox, filedialog
import sqlite3
from datetime import datetime, timedelta
import os
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

class BibliotecaApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Controle de Estoque de Livros - Biblioteca Escolar")
        self.root.geometry("1400x700")
        self.root.resizable(True, True)
        
        # Configurar ícone e tema
        self.root.configure(bg="#f0f0f0")
        
        # Inicializar banco de dados
        self.db_path = Path.home() / "biblioteca_estoque.db"
        self.inicializar_banco_dados()
        
        # Criar interface
        self.criar_interface()
        
        # Carregar dados
        self.carregar_livros()
        self.carregar_devolucoes()
        
    def inicializar_banco_dados(self):
        """Criar banco de dados SQLite se não existir"""
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        
        # Tabela de livros
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS livros (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                codigo TEXT UNIQUE NOT NULL,
                titulo TEXT NOT NULL,
                autor TEXT NOT NULL,
                editora TEXT NOT NULL,
                data_criacao TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        ''')
        
        # Tabela de empréstimos (com coluna de turma)
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS emprestimos (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                livro_id INTEGER NOT NULL,
                aluno TEXT NOT NULL,
                turma TEXT NOT NULL,
                data_retirada DATE NOT NULL,
                data_entrega_prevista DATE NOT NULL,
                data_entrega_real DATE,
                status TEXT DEFAULT 'EMPRESTADO',
                FOREIGN KEY (livro_id) REFERENCES livros(id)
            )
        ''')
        
        # Verificar e adicionar coluna turma se nao existir
        try:
            cursor.execute("PRAGMA table_info(emprestimos)")
            colunas = [coluna[1] for coluna in cursor.fetchall()]
            
            if 'turma' not in colunas:
                cursor.execute("ALTER TABLE emprestimos ADD COLUMN turma TEXT DEFAULT 'N/A'")
                print("Coluna turma adicionada ao banco de dados existente")
        except Exception as e:
            print(f"Erro ao verificar/adicionar coluna turma: {e}")
        
        conn.commit()
        conn.close()
    
    def formatar_data_br(self, data_str):
        """Converter data de YYYY-MM-DD para DD/MM/YYYY"""
        if not data_str or data_str == "None":
            return ""
        try:
            data_obj = datetime.strptime(data_str, "%Y-%m-%d")
            return data_obj.strftime("%d/%m/%Y")
        except:
            return data_str
    
    def converter_data_br_para_iso(self, data_br):
        """Converter data de DD/MM/YYYY para YYYY-MM-DD"""
        if not data_br:
            return ""
        try:
            data_obj = datetime.strptime(data_br, "%d/%m/%Y")
            return data_obj.strftime("%Y-%m-%d")
        except:
            return data_br
    
    def verificar_atraso(self, data_entrega_prevista):
        """Verificar se o livro está atrasado"""
        try:
            data_entrega = datetime.strptime(data_entrega_prevista, "%Y-%m-%d")
            hoje = datetime.now()
            return data_entrega < hoje
        except:
            return False
    
    def criar_interface(self):
        """Criar interface gráfica"""
        # Frame superior com informações da escola
        frame_header = tk.Frame(self.root, bg="#1E3C72", height=80)
        frame_header.pack(fill=tk.X, padx=0, pady=0)
        frame_header.pack_propagate(False)
        
        title_label = tk.Label(
            frame_header,
            text="📚 CONTROLE DE ESTOQUE DE LIVROS - BIBLIOTECA ESCOLAR",
            font=("Arial", 16, "bold"),
            bg="#1E3C72",
            fg="white"
        )
        title_label.pack(pady=5)
        
        school_label = tk.Label(
            frame_header,
            text="ESCOLA ESTADUAL DE EDUCAÇÃO BÁSICA NICOLAU MÜSSNICH",
            font=("Arial", 11, "bold"),
            bg="#1E3C72",
            fg="white"
        )
        school_label.pack()
        
        address_label = tk.Label(
            frame_header,
            text="Rua Professora Nely Muller 72 - Boa União - Estrela/RS",
            font=("Arial", 10),
            bg="#1E3C72",
            fg="white"
        )
        address_label.pack()
        
        # Frame para abas
        notebook = ttk.Notebook(self.root)
        notebook.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        # Aba 1: Gerenciar Livros
        self.frame_livros = ttk.Frame(notebook)
        notebook.add(self.frame_livros, text="📖 Gerenciar Livros")
        self.criar_aba_livros()
        
        # Aba 2: Empréstimos
        self.frame_emprestimos = ttk.Frame(notebook)
        notebook.add(self.frame_emprestimos, text="📤 Empréstimos")
        self.criar_aba_emprestimos()
        
        # Aba 3: Devoluções
        self.frame_devolucoes = ttk.Frame(notebook)
        notebook.add(self.frame_devolucoes, text="📥 Devoluções")
        self.criar_aba_devolucoes()
        
        # Aba 4: Relatórios
        self.frame_relatorios = ttk.Frame(notebook)
        notebook.add(self.frame_relatorios, text="📊 Relatórios")
        self.criar_aba_relatorios()
    
    def criar_aba_livros(self):
        """Criar aba para gerenciar livros"""
        # Frame de entrada
        frame_entrada = ttk.LabelFrame(self.frame_livros, text="Adicionar Novo Livro", padding=10)
        frame_entrada.pack(fill=tk.X, padx=10, pady=10)
        
        # Código
        ttk.Label(frame_entrada, text="Código do Livro:").grid(row=0, column=0, sticky=tk.W, padx=5, pady=5)
        self.entrada_codigo = ttk.Entry(frame_entrada, width=15)
        self.entrada_codigo.grid(row=0, column=1, sticky=tk.W, padx=5, pady=5)
        
        # Título
        ttk.Label(frame_entrada, text="Título:").grid(row=0, column=2, sticky=tk.W, padx=5, pady=5)
        self.entrada_titulo = ttk.Entry(frame_entrada, width=30)
        self.entrada_titulo.grid(row=0, column=3, sticky=tk.W, padx=5, pady=5)
        
        # Autor
        ttk.Label(frame_entrada, text="Autor:").grid(row=1, column=0, sticky=tk.W, padx=5, pady=5)
        self.entrada_autor = ttk.Entry(frame_entrada, width=25)
        self.entrada_autor.grid(row=1, column=1, sticky=tk.W, padx=5, pady=5)
        
        # Editora
        ttk.Label(frame_entrada, text="Editora:").grid(row=1, column=2, sticky=tk.W, padx=5, pady=5)
        self.entrada_editora = ttk.Entry(frame_entrada, width=25)
        self.entrada_editora.grid(row=1, column=3, sticky=tk.W, padx=5, pady=5)
        
        # Botão adicionar
        btn_adicionar = ttk.Button(frame_entrada, text="➕ Adicionar Livro", command=self.adicionar_livro)
        btn_adicionar.grid(row=2, column=0, columnspan=4, sticky=tk.W, padx=5, pady=10)
        
        # Frame da tabela
        frame_tabela = ttk.LabelFrame(self.frame_livros, text="Livros Cadastrados", padding=10)
        frame_tabela.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        # Treeview para livros
        self.tree_livros = ttk.Treeview(
            frame_tabela,
            columns=("Código", "Título", "Autor", "Editora", "Data Cadastro"),
            height=15,
            show="headings"
        )
        
        # Definir colunas
        self.tree_livros.column("Código", width=80)
        self.tree_livros.column("Título", width=300)
        self.tree_livros.column("Autor", width=150)
        self.tree_livros.column("Editora", width=150)
        self.tree_livros.column("Data Cadastro", width=120)
        
        # Cabeçalhos
        self.tree_livros.heading("Código", text="Código")
        self.tree_livros.heading("Título", text="Título")
        self.tree_livros.heading("Autor", text="Autor")
        self.tree_livros.heading("Editora", text="Editora")
        self.tree_livros.heading("Data Cadastro", text="Data Cadastro")
        
        # Scrollbar
        scrollbar = ttk.Scrollbar(frame_tabela, orient=tk.VERTICAL, command=self.tree_livros.yview)
        self.tree_livros.configure(yscroll=scrollbar.set)
        
        self.tree_livros.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        # Criar menu de contexto (clique direito)
        self.menu_contexto_livros = tk.Menu(self.tree_livros, tearoff=0)
        self.menu_contexto_livros.add_command(label="🗑️ Deletar Livro", command=self.deletar_livro)
        self.menu_contexto_livros.add_separator()
        self.menu_contexto_livros.add_command(label="Cancelar", command=lambda: None)
        
        # Vincular clique direito
        self.tree_livros.bind("<Button-3>", self.mostrar_menu_contexto_livros)
        
        # Botão deletar
        frame_botoes = ttk.Frame(self.frame_livros)
        frame_botoes.pack(fill=tk.X, padx=10, pady=10)
        
        btn_deletar = ttk.Button(
            frame_botoes,
            text="🗑️ Deletar Livro Selecionado",
            command=self.deletar_livro
        )
        btn_deletar.pack(side=tk.LEFT, padx=5)
        
        # Botão com atalho de teclado (Delete)
        btn_info = ttk.Label(
            frame_botoes,
            text="💡 Dica: Clique com botão direito no livro ou pressione Delete",
            foreground="#666"
        )
        btn_info.pack(side=tk.LEFT, padx=20)
        
        # Vincular tecla Delete
        self.frame_livros.bind("<Delete>", lambda e: self.deletar_livro())
    
    def criar_aba_emprestimos(self):
        """Criar aba para registrar empréstimos"""
        # Frame de entrada
        frame_entrada = ttk.LabelFrame(self.frame_emprestimos, text="Registrar Empréstimo", padding=10)
        frame_entrada.pack(fill=tk.X, padx=10, pady=10)
        
        # Livro
        ttk.Label(frame_entrada, text="Código do Livro:").grid(row=0, column=0, sticky=tk.W, padx=5, pady=5)
        self.entrada_livro_emp = ttk.Entry(frame_entrada, width=15)
        self.entrada_livro_emp.grid(row=0, column=1, sticky=tk.W, padx=5, pady=5)
        
        # Aluno
        ttk.Label(frame_entrada, text="Nome do Aluno:").grid(row=0, column=2, sticky=tk.W, padx=5, pady=5)
        self.entrada_aluno = ttk.Entry(frame_entrada, width=25)
        self.entrada_aluno.grid(row=0, column=3, sticky=tk.W, padx=5, pady=5)
        
        # Turma (NOVO)
        ttk.Label(frame_entrada, text="Turma:").grid(row=0, column=4, sticky=tk.W, padx=5, pady=5)
        self.entrada_turma = ttk.Entry(frame_entrada, width=15)
        self.entrada_turma.grid(row=0, column=5, sticky=tk.W, padx=5, pady=5)
        
        # Data de retirada
        ttk.Label(frame_entrada, text="Data de Retirada (DD/MM/YYYY):").grid(row=1, column=0, sticky=tk.W, padx=5, pady=5)
        self.entrada_data_retirada = ttk.Entry(frame_entrada, width=15)
        self.entrada_data_retirada.insert(0, datetime.now().strftime("%d/%m/%Y"))
        self.entrada_data_retirada.grid(row=1, column=1, sticky=tk.W, padx=5, pady=5)
        
        # Data de entrega prevista
        ttk.Label(frame_entrada, text="Data de Entrega Prevista (DD/MM/YYYY):").grid(row=1, column=2, sticky=tk.W, padx=5, pady=5)
        self.entrada_data_entrega = ttk.Entry(frame_entrada, width=15)
        data_entrega_padrao = (datetime.now() + timedelta(days=14)).strftime("%d/%m/%Y")
        self.entrada_data_entrega.insert(0, data_entrega_padrao)
        self.entrada_data_entrega.grid(row=1, column=3, columnspan=3, sticky=tk.W, padx=5, pady=5)
        
        # Botão registrar
        btn_registrar = ttk.Button(frame_entrada, text="📤 Registrar Empréstimo", command=self.registrar_emprestimo)
        btn_registrar.grid(row=2, column=0, columnspan=6, sticky=tk.W, padx=5, pady=10)
        
        # Frame da tabela
        frame_tabela = ttk.LabelFrame(self.frame_emprestimos, text="Empréstimos Ativos", padding=10)
        frame_tabela.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        # Treeview para empréstimos
        self.tree_emprestimos = ttk.Treeview(
            frame_tabela,
            columns=("Código", "Título", "Aluno", "Turma", "Retirada", "Entrega Prevista", "Status"),
            height=15,
            show="headings"
        )
        
        # Definir colunas
        self.tree_emprestimos.column("Código", width=70)
        self.tree_emprestimos.column("Título", width=220)
        self.tree_emprestimos.column("Aluno", width=100)
        self.tree_emprestimos.column("Turma", width=70)
        self.tree_emprestimos.column("Retirada", width=90)
        self.tree_emprestimos.column("Entrega Prevista", width=110)
        self.tree_emprestimos.column("Status", width=120)
        
        # Cabeçalhos
        self.tree_emprestimos.heading("Código", text="Código")
        self.tree_emprestimos.heading("Título", text="Título")
        self.tree_emprestimos.heading("Aluno", text="Aluno")
        self.tree_emprestimos.heading("Turma", text="Turma")
        self.tree_emprestimos.heading("Retirada", text="Retirada")
        self.tree_emprestimos.heading("Entrega Prevista", text="Entrega Prevista")
        self.tree_emprestimos.heading("Status", text="Status")
        
        # Scrollbar
        scrollbar = ttk.Scrollbar(frame_tabela, orient=tk.VERTICAL, command=self.tree_emprestimos.yview)
        self.tree_emprestimos.configure(yscroll=scrollbar.set)
        
        self.tree_emprestimos.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        # Configurar estilo para linhas atrasadas
        style = ttk.Style()
        style.configure("Treeview", rowheight=25)
    
    def criar_aba_devolucoes(self):
        """Criar aba para registrar devoluções"""
        # Frame de entrada
        frame_entrada = ttk.LabelFrame(self.frame_devolucoes, text="Registrar Devolução", padding=10)
        frame_entrada.pack(fill=tk.X, padx=10, pady=10)
        
        # ID do empréstimo
        ttk.Label(frame_entrada, text="Selecione o empréstimo abaixo e clique em Devolver").grid(row=0, column=0, columnspan=4, sticky=tk.W, padx=5, pady=5)
        
        # Botão devolver
        btn_devolver = ttk.Button(frame_entrada, text="✓ Marcar como Devolvido", command=self.marcar_devolucao)
        btn_devolver.grid(row=1, column=0, columnspan=4, sticky=tk.W, padx=5, pady=10)
        
        # Frame da tabela
        frame_tabela = ttk.LabelFrame(self.frame_devolucoes, text="Empréstimos Pendentes de Devolução", padding=10)
        frame_tabela.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        # Treeview para devoluções
        self.tree_devolucoes = ttk.Treeview(
            frame_tabela,
            columns=("ID", "Código", "Título", "Aluno", "Turma", "Retirada", "Entrega Prevista", "Status"),
            height=15,
            show="headings"
        )
        
        # Definir colunas
        self.tree_devolucoes.column("ID", width=40)
        self.tree_devolucoes.column("Código", width=70)
        self.tree_devolucoes.column("Título", width=200)
        self.tree_devolucoes.column("Aluno", width=100)
        self.tree_devolucoes.column("Turma", width=70)
        self.tree_devolucoes.column("Retirada", width=90)
        self.tree_devolucoes.column("Entrega Prevista", width=110)
        self.tree_devolucoes.column("Status", width=100)
        
        # Cabeçalhos
        self.tree_devolucoes.heading("ID", text="ID")
        self.tree_devolucoes.heading("Código", text="Código")
        self.tree_devolucoes.heading("Título", text="Título")
        self.tree_devolucoes.heading("Aluno", text="Aluno")
        self.tree_devolucoes.heading("Turma", text="Turma")
        self.tree_devolucoes.heading("Retirada", text="Retirada")
        self.tree_devolucoes.heading("Entrega Prevista", text="Entrega Prevista")
        self.tree_devolucoes.heading("Status", text="Status")
        
        # Scrollbar
        scrollbar = ttk.Scrollbar(frame_tabela, orient=tk.VERTICAL, command=self.tree_devolucoes.yview)
        self.tree_devolucoes.configure(yscroll=scrollbar.set)
        
        self.tree_devolucoes.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
    
    def criar_aba_relatorios(self):
        """Criar aba para relatórios"""
        # Frame de filtros
        frame_filtros = ttk.LabelFrame(self.frame_relatorios, text="Filtros", padding=10)
        frame_filtros.pack(fill=tk.X, padx=10, pady=10)
        
        btn_atrasos = ttk.Button(
            frame_filtros,
            text="📋 Livros Atrasados",
            command=self.mostrar_atrasos
        )
        btn_atrasos.pack(side=tk.LEFT, padx=5)
        
        btn_todos = ttk.Button(
            frame_filtros,
            text="📋 Todos os Empréstimos",
            command=self.mostrar_todos_emprestimos
        )
        btn_todos.pack(side=tk.LEFT, padx=5)
        
        btn_devolvidos = ttk.Button(
            frame_filtros,
            text="📋 Livros Devolvidos",
            command=self.mostrar_devolvidos
        )
        btn_devolvidos.pack(side=tk.LEFT, padx=5)
        
        btn_exportar = ttk.Button(
            frame_filtros,
            text="💾 Exportar para Word",
            command=self.exportar_word
        )
        btn_exportar.pack(side=tk.LEFT, padx=5)
        
        # Frame da tabela
        frame_tabela = ttk.LabelFrame(self.frame_relatorios, text="Relatório", padding=10)
        frame_tabela.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        # Treeview para relatórios
        self.tree_relatorios = ttk.Treeview(
            frame_tabela,
            columns=("ID", "Código", "Título", "Aluno", "Turma", "Retirada", "Entrega Prevista", "Entrega Real", "Status"),
            height=15,
            show="headings"
        )
        
        # Definir colunas
        self.tree_relatorios.column("ID", width=40)
        self.tree_relatorios.column("Código", width=70)
        self.tree_relatorios.column("Título", width=180)
        self.tree_relatorios.column("Aluno", width=90)
        self.tree_relatorios.column("Turma", width=70)
        self.tree_relatorios.column("Retirada", width=90)
        self.tree_relatorios.column("Entrega Prevista", width=110)
        self.tree_relatorios.column("Entrega Real", width=110)
        self.tree_relatorios.column("Status", width=100)
        
        # Cabeçalhos
        self.tree_relatorios.heading("ID", text="ID")
        self.tree_relatorios.heading("Código", text="Código")
        self.tree_relatorios.heading("Título", text="Título")
        self.tree_relatorios.heading("Aluno", text="Aluno")
        self.tree_relatorios.heading("Turma", text="Turma")
        self.tree_relatorios.heading("Retirada", text="Retirada")
        self.tree_relatorios.heading("Entrega Prevista", text="Entrega Prevista")
        self.tree_relatorios.heading("Entrega Real", text="Entrega Real")
        self.tree_relatorios.heading("Status", text="Status")
        
        # Scrollbar
        scrollbar = ttk.Scrollbar(frame_tabela, orient=tk.VERTICAL, command=self.tree_relatorios.yview)
        self.tree_relatorios.configure(yscroll=scrollbar.set)
        
        self.tree_relatorios.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
    
    def adicionar_livro(self):
        """Adicionar novo livro"""
        codigo = self.entrada_codigo.get().strip()
        titulo = self.entrada_titulo.get().strip()
        autor = self.entrada_autor.get().strip()
        editora = self.entrada_editora.get().strip()
        
        if not all([codigo, titulo, autor, editora]):
            messagebox.showerror("Erro", "Preencha todos os campos!")
            return
        
        try:
            conn = sqlite3.connect(self.db_path)
            cursor = conn.cursor()
            cursor.execute(
                "INSERT INTO livros (codigo, titulo, autor, editora) VALUES (?, ?, ?, ?)",
                (codigo, titulo, autor, editora)
            )
            conn.commit()
            conn.close()
            
            messagebox.showinfo("Sucesso", "Livro adicionado com sucesso!")
            
            # Limpar campos
            self.entrada_codigo.delete(0, tk.END)
            self.entrada_titulo.delete(0, tk.END)
            self.entrada_autor.delete(0, tk.END)
            self.entrada_editora.delete(0, tk.END)
            
            # Recarregar tabela
            self.carregar_livros()
        except sqlite3.IntegrityError:
            messagebox.showerror("Erro", "Código de livro já existe!")
        except Exception as e:
            messagebox.showerror("Erro", f"Erro ao adicionar livro: {str(e)}")
    
    def mostrar_menu_contexto_livros(self, event):
        """Mostrar menu de contexto ao clicar com botao direito"""
        # Selecionar o item sob o cursor
        item = self.tree_livros.identify_row(event.y)
        if item:
            self.tree_livros.selection_set(item)
        
        # Mostrar menu
        try:
            self.menu_contexto_livros.tk_popup(event.x_root, event.y_root)
        finally:
            self.menu_contexto_livros.grab_release()
    
    def deletar_livro(self):
        """Deletar livro selecionado"""
        selecionado = self.tree_livros.selection()
        if not selecionado:
            messagebox.showwarning("Aviso", "Selecione um livro para deletar!")
            return
        
        item = self.tree_livros.item(selecionado[0])
        codigo = item['values'][0]
        
        if messagebox.askyesno("Confirmar", f"Deletar livro com código {codigo}?"):
            try:
                conn = sqlite3.connect(self.db_path)
                cursor = conn.cursor()
                cursor.execute("DELETE FROM livros WHERE codigo = ?", (codigo,))
                conn.commit()
                conn.close()
                
                messagebox.showinfo("Sucesso", "Livro deletado com sucesso!")
                self.carregar_livros()
            except Exception as e:
                messagebox.showerror("Erro", f"Erro ao deletar livro: {str(e)}")
    
    def carregar_livros(self):
        """Carregar livros na tabela"""
        # Limpar tabela
        for item in self.tree_livros.get_children():
            self.tree_livros.delete(item)
        
        try:
            conn = sqlite3.connect(self.db_path)
            cursor = conn.cursor()
            cursor.execute("SELECT codigo, titulo, autor, editora, data_criacao FROM livros ORDER BY titulo")
            livros = cursor.fetchall()
            conn.close()
            
            for livro in livros:
                data_formatada = self.formatar_data_br(livro[4])
                self.tree_livros.insert("", tk.END, values=(livro[0], livro[1], livro[2], livro[3], data_formatada))
        except Exception as e:
            messagebox.showerror("Erro", f"Erro ao carregar livros: {str(e)}")
    
    def registrar_emprestimo(self):
        """Registrar novo empréstimo"""
        codigo_livro = self.entrada_livro_emp.get().strip()
        aluno = self.entrada_aluno.get().strip()
        turma = self.entrada_turma.get().strip()
        data_retirada = self.entrada_data_retirada.get().strip()
        data_entrega = self.entrada_data_entrega.get().strip()
        
        if not all([codigo_livro, aluno, turma, data_retirada, data_entrega]):
            messagebox.showerror("Erro", "Preencha todos os campos!")
            return
        
        try:
            # Validar e converter datas
            data_retirada_iso = self.converter_data_br_para_iso(data_retirada)
            data_entrega_iso = self.converter_data_br_para_iso(data_entrega)
            
            datetime.strptime(data_retirada_iso, "%Y-%m-%d")
            datetime.strptime(data_entrega_iso, "%Y-%m-%d")
            
            conn = sqlite3.connect(self.db_path)
            cursor = conn.cursor()
            
            # Verificar se livro existe
            cursor.execute("SELECT id FROM livros WHERE codigo = ?", (codigo_livro,))
            livro = cursor.fetchone()
            
            if not livro:
                messagebox.showerror("Erro", "Livro não encontrado!")
                conn.close()
                return
            
            livro_id = livro[0]
            
            # Registrar empréstimo
            cursor.execute(
                "INSERT INTO emprestimos (livro_id, aluno, turma, data_retirada, data_entrega_prevista, status) VALUES (?, ?, ?, ?, ?, ?)",
                (livro_id, aluno, turma, data_retirada_iso, data_entrega_iso, "EMPRESTADO")
            )
            conn.commit()
            conn.close()
            
            messagebox.showinfo("Sucesso", "Empréstimo registrado com sucesso!")
            
            # Limpar campos
            self.entrada_livro_emp.delete(0, tk.END)
            self.entrada_aluno.delete(0, tk.END)
            self.entrada_turma.delete(0, tk.END)
            self.entrada_data_retirada.delete(0, tk.END)
            self.entrada_data_retirada.insert(0, datetime.now().strftime("%d/%m/%Y"))
            self.entrada_data_entrega.delete(0, tk.END)
            data_entrega_padrao = (datetime.now() + timedelta(days=14)).strftime("%d/%m/%Y")
            self.entrada_data_entrega.insert(0, data_entrega_padrao)
            
            # Recarregar tabelas
            self.carregar_emprestimos()
            self.carregar_devolucoes()
        except ValueError:
            messagebox.showerror("Erro", "Formato de data inválido! Use DD/MM/YYYY")
        except Exception as e:
            messagebox.showerror("Erro", f"Erro ao registrar empréstimo: {str(e)}")
    
    def carregar_emprestimos(self):
        """Carregar empréstimos ativos na tabela"""
        # Limpar tabela
        for item in self.tree_emprestimos.get_children():
            self.tree_emprestimos.delete(item)
        
        try:
            conn = sqlite3.connect(self.db_path)
            cursor = conn.cursor()
            cursor.execute("""
                SELECT e.id, l.codigo, l.titulo, e.aluno, e.turma, e.data_retirada, 
                       e.data_entrega_prevista, e.status
                FROM emprestimos e
                JOIN livros l ON e.livro_id = l.id
                WHERE e.status = 'EMPRESTADO'
                ORDER BY e.data_entrega_prevista
            """)
            emprestimos = cursor.fetchall()
            conn.close()
            
            for emp in emprestimos:
                data_retirada_br = self.formatar_data_br(emp[5])
                data_entrega_br = self.formatar_data_br(emp[6])
                
                # Verificar se está atrasado
                if self.verificar_atraso(emp[6]):
                    status = "🔴 ATRASADO"
                else:
                    status = "🟢 NO PRAZO"
                
                valores = (emp[1], emp[2], emp[3], emp[4], data_retirada_br, data_entrega_br, status)
                item_id = self.tree_emprestimos.insert("", tk.END, values=valores)
                
                # Colorir linha se atrasado
                if status == "🔴 ATRASADO":
                    self.tree_emprestimos.item(item_id, tags=('atrasado',))
            
            # Configurar tag para linhas atrasadas
            self.tree_emprestimos.tag_configure('atrasado', background='#ffcccc', foreground='#cc0000')
        except Exception as e:
            messagebox.showerror("Erro", f"Erro ao carregar empréstimos: {str(e)}")
    
    def carregar_devolucoes(self):
        """Carregar empréstimos pendentes de devolução"""
        # Limpar tabela
        for item in self.tree_devolucoes.get_children():
            self.tree_devolucoes.delete(item)
        
        try:
            conn = sqlite3.connect(self.db_path)
            cursor = conn.cursor()
            cursor.execute("""
                SELECT e.id, l.codigo, l.titulo, e.aluno, e.turma, e.data_retirada, 
                       e.data_entrega_prevista, e.status
                FROM emprestimos e
                JOIN livros l ON e.livro_id = l.id
                WHERE e.status = 'EMPRESTADO'
                ORDER BY e.data_entrega_prevista
            """)
            emprestimos = cursor.fetchall()
            conn.close()
            
            if not emprestimos:
                item_id = self.tree_devolucoes.insert("", tk.END, values=("", "", "Nenhum emprestimo pendente", "", "", "", "", "Tudo devolvido!"))
                self.tree_devolucoes.item(item_id, tags=('completo',))
                self.tree_devolucoes.tag_configure('completo', background='#ccffcc', foreground='#00aa00')
                return
            
            for emp in emprestimos:
                data_retirada_br = self.formatar_data_br(emp[5])
                data_entrega_br = self.formatar_data_br(emp[6])
                
                # Verificar se está atrasado
                if self.verificar_atraso(emp[6]):
                    status = "🔴 ATRASADO"
                else:
                    status = "🟢 NO PRAZO"
                
                item_id = self.tree_devolucoes.insert("", tk.END, values=(emp[0], emp[1], emp[2], emp[3], emp[4], data_retirada_br, data_entrega_br, status))
                
                # Colorir linha se atrasado
                if status == "🔴 ATRASADO":
                    self.tree_devolucoes.item(item_id, tags=('atrasado',))
            
            # Configurar tag para linhas atrasadas
            self.tree_devolucoes.tag_configure('atrasado', background='#ffcccc', foreground='#cc0000')
        except Exception as e:
            messagebox.showerror("Erro", f"Erro ao carregar devoluções: {str(e)}")
    
    def marcar_devolucao(self):
        """Marcar empréstimo como devolvido"""
        selecionado = self.tree_devolucoes.selection()
        if not selecionado:
            messagebox.showwarning("Aviso", "Selecione um empréstimo para marcar como devolvido!")
            return
        
        item = self.tree_devolucoes.item(selecionado[0])
        emp_id = item['values'][0]
        
        try:
            conn = sqlite3.connect(self.db_path)
            cursor = conn.cursor()
            cursor.execute(
                "UPDATE emprestimos SET status = 'DEVOLVIDO', data_entrega_real = ? WHERE id = ?",
                (datetime.now().strftime("%Y-%m-%d"), emp_id)
            )
            conn.commit()
            conn.close()
            
            messagebox.showinfo("Sucesso", "Livro marcado como devolvido!")
            self.carregar_devolucoes()
            self.carregar_emprestimos()
        except Exception as e:
            messagebox.showerror("Erro", f"Erro ao marcar devolução: {str(e)}")
    
    def mostrar_atrasos(self):
        """Mostrar apenas livros atrasados"""
        # Limpar tabela
        for item in self.tree_relatorios.get_children():
            self.tree_relatorios.delete(item)
        
        try:
            conn = sqlite3.connect(self.db_path)
            cursor = conn.cursor()
            cursor.execute("""
                SELECT e.id, l.codigo, l.titulo, e.aluno, e.turma, e.data_retirada, 
                       e.data_entrega_prevista, e.data_entrega_real, e.status
                FROM emprestimos e
                JOIN livros l ON e.livro_id = l.id
                WHERE e.status = 'EMPRESTADO' AND date(e.data_entrega_prevista) < date('now')
                ORDER BY e.data_entrega_prevista
            """)
            emprestimos = cursor.fetchall()
            conn.close()
            
            if not emprestimos:
                messagebox.showinfo("Relatório", "Nenhum livro atrasado!")
                return
            
            for emp in emprestimos:
                data_retirada_br = self.formatar_data_br(emp[5])
                data_entrega_br = self.formatar_data_br(emp[6])
                data_entrega_real_br = self.formatar_data_br(emp[7])
                
                item_id = self.tree_relatorios.insert("", tk.END, values=(emp[0], emp[1], emp[2], emp[3], emp[4], data_retirada_br, data_entrega_br, data_entrega_real_br, "🔴 ATRASADO"))
                self.tree_relatorios.item(item_id, tags=('atrasado',))
            
            # Configurar tag para linhas atrasadas
            self.tree_relatorios.tag_configure('atrasado', background='#ffcccc', foreground='#cc0000')
        except Exception as e:
            messagebox.showerror("Erro", f"Erro ao gerar relatório: {str(e)}")
    
    def mostrar_todos_emprestimos(self):
        """Mostrar todos os empréstimos"""
        # Limpar tabela
        for item in self.tree_relatorios.get_children():
            self.tree_relatorios.delete(item)
        
        try:
            conn = sqlite3.connect(self.db_path)
            cursor = conn.cursor()
            cursor.execute("""
                SELECT e.id, l.codigo, l.titulo, e.aluno, e.turma, e.data_retirada, 
                       e.data_entrega_prevista, e.data_entrega_real, e.status
                FROM emprestimos e
                JOIN livros l ON e.livro_id = l.id
                ORDER BY e.data_retirada DESC
            """)
            emprestimos = cursor.fetchall()
            conn.close()
            
            for emp in emprestimos:
                data_retirada_br = self.formatar_data_br(emp[5])
                data_entrega_br = self.formatar_data_br(emp[6])
                data_entrega_real_br = self.formatar_data_br(emp[7])
                
                # Verificar se está atrasado
                if emp[8] == 'EMPRESTADO' and self.verificar_atraso(emp[6]):
                    status = "🔴 ATRASADO"
                    item_id = self.tree_relatorios.insert("", tk.END, values=(emp[0], emp[1], emp[2], emp[3], emp[4], data_retirada_br, data_entrega_br, data_entrega_real_br, status))
                    self.tree_relatorios.item(item_id, tags=('atrasado',))
                else:
                    status = emp[8]
                    self.tree_relatorios.insert("", tk.END, values=(emp[0], emp[1], emp[2], emp[3], emp[4], data_retirada_br, data_entrega_br, data_entrega_real_br, status))
            
            # Configurar tag para linhas atrasadas
            self.tree_relatorios.tag_configure('atrasado', background='#ffcccc', foreground='#cc0000')
        except Exception as e:
            messagebox.showerror("Erro", f"Erro ao gerar relatório: {str(e)}")
    
    def mostrar_devolvidos(self):
        """Mostrar apenas livros devolvidos"""
        # Limpar tabela
        for item in self.tree_relatorios.get_children():
            self.tree_relatorios.delete(item)
        
        try:
            conn = sqlite3.connect(self.db_path)
            cursor = conn.cursor()
            cursor.execute("""
                SELECT e.id, l.codigo, l.titulo, e.aluno, e.turma, e.data_retirada, 
                       e.data_entrega_prevista, e.data_entrega_real, e.status
                FROM emprestimos e
                JOIN livros l ON e.livro_id = l.id
                WHERE e.status = 'DEVOLVIDO'
                ORDER BY e.data_entrega_real DESC
            """)
            emprestimos = cursor.fetchall()
            conn.close()
            
            for emp in emprestimos:
                data_retirada_br = self.formatar_data_br(emp[5])
                data_entrega_br = self.formatar_data_br(emp[6])
                data_entrega_real_br = self.formatar_data_br(emp[7])
                self.tree_relatorios.insert("", tk.END, values=(emp[0], emp[1], emp[2], emp[3], emp[4], data_retirada_br, data_entrega_br, data_entrega_real_br, emp[8]))
        except Exception as e:
            messagebox.showerror("Erro", f"Erro ao gerar relatório: {str(e)}")
    
    def exportar_word(self):
        """Exportar relatório para Word (.docx)"""
        arquivo = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word files", "*.docx"), ("All files", "*.*")]
        )
        
        if not arquivo:
            return
        
        try:
            # Criar documento Word
            doc = Document()
            
            # Adicionar cabeçalho
            heading = doc.add_heading("RELATÓRIO DE CONTROLE DE ESTOQUE DE LIVROS", 0)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            school_para = doc.add_paragraph("ESCOLA ESTADUAL DE EDUCAÇÃO BÁSICA NICOLAU MÜSSNICH")
            school_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            school_para.runs[0].font.bold = True
            
            address_para = doc.add_paragraph("Rua Professora Nely Muller 72 - Boa União - Estrela/RS")
            address_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            date_para = doc.add_paragraph(f"Data do Relatório: {datetime.now().strftime('%d/%m/%Y às %H:%M')}")
            date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.add_paragraph()  # Linha em branco
            
            # Obter dados da tabela
            conn = sqlite3.connect(self.db_path)
            cursor = conn.cursor()
            cursor.execute("""
                SELECT e.id, l.codigo, l.titulo, e.aluno, e.turma, e.data_retirada, 
                       e.data_entrega_prevista, e.data_entrega_real, e.status
                FROM emprestimos e
                JOIN livros l ON e.livro_id = l.id
                ORDER BY e.data_retirada DESC
            """)
            emprestimos = cursor.fetchall()
            conn.close()
            
            if not emprestimos:
                doc.add_paragraph("Nenhum registro encontrado.")
            else:
                # Criar tabela
                table = doc.add_table(rows=1, cols=9)
                table.style = 'Light Grid Accent 1'
                
                # Cabeçalho da tabela
                hdr_cells = table.rows[0].cells
                headers = ["ID", "Código", "Título", "Aluno", "Turma", "Data Retirada", "Data Entrega Prevista", "Data Entrega Real", "Status"]
                for i, header in enumerate(headers):
                    hdr_cells[i].text = header
                    # Formatar cabeçalho
                    for paragraph in hdr_cells[i].paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
                
                # Adicionar dados
                for emp in emprestimos:
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(emp[0])
                    row_cells[1].text = str(emp[1])
                    row_cells[2].text = str(emp[2])
                    row_cells[3].text = str(emp[3])
                    row_cells[4].text = str(emp[4])
                    row_cells[5].text = self.formatar_data_br(emp[5])
                    row_cells[6].text = self.formatar_data_br(emp[6])
                    row_cells[7].text = self.formatar_data_br(emp[7])
                    
                    # Verificar se está atrasado
                    if emp[8] == 'EMPRESTADO' and self.verificar_atraso(emp[6]):
                        row_cells[8].text = "🔴 ATRASADO"
                    else:
                        row_cells[8].text = str(emp[8])
            
            # Salvar documento
            doc.save(arquivo)
            messagebox.showinfo("Sucesso", f"Relatório exportado para {arquivo}")
        except Exception as e:
            messagebox.showerror("Erro", f"Erro ao exportar: {str(e)}")


def main():
    root = tk.Tk()
    app = BibliotecaApp(root)
    root.mainloop()


if __name__ == "__main__":
    main()
